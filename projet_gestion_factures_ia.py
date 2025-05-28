#!/usr/bin/env python3
"""
Générateur de document de projet professionnel - FacturIQ.ai
===========================================================

Génère un document PDF et Word professionnel décrivant le projet
"FacturIQ.ai - La Facturation Intelligente" selon le template requis.
"""

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    Image,
    PageBreak,
)
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from datetime import datetime
import os

# Pour la génération Word
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ Module python-docx non disponible. Installation: pip install python-docx")


class ProjetDocumentGenerator:
    """Générateur de document de projet professionnel."""

    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()

    def _setup_custom_styles(self):
        """Configuration des styles personnalisés avec Times New Roman."""
        # Style pour le titre principal
        self.title_style = ParagraphStyle(
            "CustomTitle",
            parent=self.styles["Title"],
            fontSize=20,  # Plus grand pour le titre
            fontName="Times-Bold",
            textColor=colors.HexColor("#2E4057"),
            spaceAfter=30,
            alignment=TA_CENTER,
            leading=30,  # 1.5 * fontSize
        )

        # Style pour les titres de section
        self.section_style = ParagraphStyle(
            "SectionTitle",
            parent=self.styles["Heading1"],
            fontSize=16,  # Taille appropriée pour les headers
            fontName="Times-Bold",
            textColor=colors.HexColor("#048A81"),
            spaceBefore=20,
            spaceAfter=12,
            borderPadding=5,
            backColor=colors.HexColor("#F0F8FF"),
            leading=24,  # 1.5 * fontSize
        )

        # Style pour les sous-titres
        self.subsection_style = ParagraphStyle(
            "SubsectionTitle",
            parent=self.styles["Heading2"],
            fontSize=14,  # Taille appropriée pour les sous-headers
            fontName="Times-Bold",
            textColor=colors.HexColor("#2E4057"),
            spaceBefore=15,
            spaceAfter=8,
            leading=21,  # 1.5 * fontSize
        )

        # Style pour le texte normal (Times New Roman 12pt, espacement 1.5)
        self.body_style = ParagraphStyle(
            "CustomBody",
            parent=self.styles["Normal"],
            fontSize=12,  # Taille 12 comme demandé
            fontName="Times-Roman",
            textColor=colors.HexColor("#333333"),
            alignment=TA_JUSTIFY,
            spaceAfter=12,
            leftIndent=0,
            rightIndent=0,
            leading=18,  # 1.5 * 12 = 18 pour espacement 1.5
        )

        # Style pour les points importants
        self.highlight_style = ParagraphStyle(
            "Highlight",
            parent=self.body_style,
            fontSize=12,
            fontName="Times-Bold",
            textColor=colors.HexColor("#048A81"),
            leftIndent=20,
            leading=18,  # 1.5 * 12
        )

    def generate_documents(self):
        """Génère les documents PDF et Word."""
        pdf_filename = self.generate_pdf_document()
        word_filename = self.generate_word_document() if DOCX_AVAILABLE else None

        return pdf_filename, word_filename

    def generate_pdf_document(self, filename="facturiq_ai.pdf"):
        """Génère le document PDF complet."""
        doc = SimpleDocTemplate(
            filename,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )

        story = []

        # Page de titre
        story.extend(self._create_title_page())

        # Table des matières
        story.append(PageBreak())
        story.extend(self._create_table_of_contents())

        # Contenu principal
        story.append(PageBreak())
        story.extend(self._create_main_content())

        # Génération du PDF
        doc.build(story)
        print(f"✅ Document PDF généré avec succès: {filename}")
        return filename

    def generate_word_document(self, filename="facturiq_ai.docx"):
        """Génère le document Word complet."""
        if not DOCX_AVAILABLE:
            print(
                "❌ Impossible de générer le document Word - module python-docx manquant"
            )
            return None

        doc = Document()

        # Configuration des styles
        self._setup_word_styles(doc)

        # Page de titre
        self._create_word_title_page(doc)

        # Saut de page
        doc.add_page_break()

        # Table des matières
        self._create_word_table_of_contents(doc)

        # Saut de page
        doc.add_page_break()

        # Contenu principal
        self._create_word_main_content(doc)

        # Sauvegarde
        doc.save(filename)
        print(f"✅ Document Word généré avec succès: {filename}")
        return filename

    def _setup_word_styles(self, doc):
        """Configure les styles Word."""
        styles = doc.styles

        # Style pour le texte normal (Times New Roman 12pt, espacement 1.5)
        try:
            normal_style = styles["Normal"]
            font = normal_style.font
            font.name = "Times New Roman"
            font.size = Pt(12)

            paragraph_format = normal_style.paragraph_format
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            paragraph_format.line_spacing = 1.5
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        except:
            pass

    def _create_word_title_page(self, doc):
        """Crée la page de titre Word."""
        # Titre principal
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("DOCUMENT DE PROJET")
        run.font.name = "Times New Roman"
        run.font.size = Pt(20)
        run.font.bold = True

        doc.add_paragraph()  # Espace

        # Nom du projet
        project_title = doc.add_paragraph()
        project_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = project_title.add_run("FacturIQ.ai\nLa Facturation Intelligente")
        run.font.name = "Times New Roman"
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = None  # Couleur par défaut

        doc.add_paragraph()  # Espace

        # Sous-titre
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(
            "Plateforme d'automatisation intelligente pour la gestion\net génération de factures professionnelles"
        )
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

        doc.add_paragraph()  # Espace
        doc.add_paragraph()  # Espace

        # Informations du document
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_text = f"""Date: {datetime.now().strftime('%d/%m/%Y')}
Version: 2.0
Statut: Production
Technologie: Azure AI Foundry + GPT-4o
Développeur: Benny Joumessi Jason Isaac"""

        run = info.add_run(info_text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    def _create_word_table_of_contents(self, doc):
        """Crée la table des matières Word."""
        # Titre de la table des matières
        toc_title = doc.add_paragraph()
        run = toc_title.add_run("TABLE DES MATIÈRES")
        run.font.name = "Times New Roman"
        run.font.size = Pt(16)
        run.font.bold = True

        doc.add_paragraph()  # Espace

        # Contenu de la table
        toc_items = [
            "1. Présentation de la start-up et du projet",
            "2. Problème à résoudre et opportunité de marché",
            "3. Facteur innovant",
            "4. Résumé du projet",
            "5. Cible et bénéficiaires",
            "6. Modèle économique",
            "7. Équipe de projet",
        ]

        for item in toc_items:
            p = doc.add_paragraph()
            run = p.add_run(item)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    def _create_word_main_content(self, doc):
        """Crée le contenu principal Word."""
        # Section 1
        self._add_word_section_title(doc, "1. PRÉSENTATION DE LA START-UP ET DU PROJET")

        self._add_word_subsection(doc, "Notre Vision")
        self._add_word_paragraph(
            doc,
            "FacturIQ.ai est une start-up technologique camerounaise spécialisée dans l'automatisation "
            "intelligente des processus de facturation pour les PME et entreprises de services. Notre mission "
            "est de révolutionner la gestion de facturation en combinant l'intelligence artificielle "
            "avancée avec les technologies cloud Microsoft Azure.",
        )

        self._add_word_subsection(doc, "Le Projet")
        self._add_word_paragraph(
            doc,
            "FacturIQ.ai - La Facturation Intelligente est une plateforme complète qui automatise "
            "entièrement le cycle de vie des factures, de la génération à l'analyse des performances "
            "commerciales. Utilisant GPT-4o d'OpenAI via Azure AI Foundry, le système comprend le "
            "langage naturel et génère des factures professionnelles en quelques secondes.",
        )

        # Section 2
        self._add_word_section_title(
            doc, "2. PROBLÈME À RÉSOUDRE ET OPPORTUNITÉ DE MARCHÉ"
        )

        self._add_word_subsection(doc, "Problématiques Identifiées")
        problems = [
            "• Processus manuel chronophage : Les PME passent en moyenne 5-8 heures par semaine sur la facturation",
            "• Erreurs humaines coûteuses : 23% des factures contiennent des erreurs, retardant les paiements",
            "• Manque de visibilité : Absence d'outils d'analyse pour optimiser la trésorerie",
            "• Conformité réglementaire : Difficultés à respecter les normes fiscales en évolution",
            "• Processus de suivi défaillant : Retards de paiement non anticipés",
        ]

        for problem in problems:
            self._add_word_highlight(doc, problem)

        self._add_word_subsection(doc, "Opportunité de Marché")
        self._add_word_paragraph(
            doc,
            "Le marché mondial des logiciels de facturation représente 18,7 milliards USD en 2024, "
            "avec une croissance annuelle de 15,4%. Au Cameroun et en Afrique centrale, les PME "
            "cherchent des solutions d'automatisation adaptées au contexte local. Notre approche "
            "IA-first nous positionne sur un segment à forte valeur ajoutée avec peu de concurrence directe.",
        )

        # Section 3
        self._add_word_section_title(doc, "3. FACTEUR INNOVANT")

        innovations = [
            '• IA Conversationnelle : Génération de factures par langage naturel ("Facture Acme Corp, 40h consulting, 75000 FCFA/h")',
            "• Optimisation Avancée : Système de cache intelligent offrant 23,464x d'amélioration des performances",
            "• Analytics Prédictive : Prévisions de trésorerie et scoring de risque client automatisés",
            "• Architecture Résiliente : Circuit breaker et rate limiting pour 99.9% de disponibilité",
            "• Interface Unifiée : Chat IA + dashboard analytics + génération rapide dans une seule plateforme",
        ]

        for innovation in innovations:
            self._add_word_highlight(doc, innovation)

        self._add_word_paragraph(
            doc,
            "Notre différenciation clé réside dans l'intégration native de l'IA générative avec "
            "des capacités d'analyse business en temps réel, créant un assistant intelligent "
            "capable de comprendre le contexte métier camerounais et d'anticiper les besoins.",
        )

        # Section 4
        self._add_word_section_title(doc, "4. RÉSUMÉ DU PROJET")

        self._add_word_subsection(doc, "Architecture Technique")
        tech_points = [
            "• Frontend : Interface Streamlit responsive avec dashboard analytics interactif",
            "• IA Core : GPT-4o via Azure AI Foundry pour génération intelligente",
            "• Backend : Microservices Python avec pattern singleton et cache avancé",
            "• Storage : CosmosDB (données), Blob Storage (PDFs), AI Search (indexation)",
            "• Analytics : Plotly pour visualisations, algorithmes prédictifs personnalisés",
        ]

        for point in tech_points:
            self._add_word_highlight(doc, point)

        self._add_word_subsection(doc, "Fonctionnalités Clés")
        self._add_word_paragraph(
            doc,
            "La plateforme offre une génération de factures en langage naturel, un système "
            "d'analytics business en temps réel, un cache intelligent réduisant les coûts API "
            "de 75%, et une architecture cloud-native garantissant scalabilité et résilience. "
            "L'interface unifiée combine chat IA, dashboard analytics et outils de gestion "
            "dans une expérience utilisateur moderne adaptée au marché camerounais.",
        )

        # Section 5
        self._add_word_section_title(doc, "5. CIBLE ET BÉNÉFICIAIRES")

        self._add_word_subsection(doc, "Marché Primaire")
        targets = [
            "• PME Services (5-50 employés) : Consultants, agences, services professionnels au Cameroun",
            "• Freelances Premium : Consultants indépendants à forte facturation",
            "• Start-ups B2B : Entreprises SaaS en croissance cherchant l'automatisation",
            "• Cabinets Spécialisés : Comptables, avocats, architectes, ingénieurs",
        ]

        for target in targets:
            self._add_word_highlight(doc, target)

        self._add_word_subsection(doc, "Bénéfices Quantifiés")
        benefits = [
            "• Gain de temps : 85% de réduction du temps de facturation",
            "• Réduction erreurs : 90% moins d'erreurs grâce à l'automatisation IA",
            "• Amélioration trésorerie : 25% d'accélération des encaissements",
            "• Visibilité business : Analytics prédictive pour optimisation continue",
        ]

        for benefit in benefits:
            self._add_word_highlight(doc, benefit)

        # Section 6
        self._add_word_section_title(doc, "6. MODÈLE ÉCONOMIQUE")

        self._add_word_subsection(doc, "Stratégie de Monétisation")

        # Table des prix
        table = doc.add_table(rows=5, cols=4)
        table.style = "Table Grid"

        # En-têtes
        headers = ["Forfait", "Prix/mois", "Factures/mois", "Fonctionnalités"]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        # Données
        pricing_data = [
            ["Starter", "15,000 FCFA", "≤ 50", "Génération IA + Analytics de base"],
            ["Professional", "40,000 FCFA", "≤ 200", "+ Dashboard avancé + API"],
            [
                "Enterprise",
                "100,000 FCFA",
                "Illimité",
                "+ IA personnalisée + Support dédié",
            ],
            ["White Label", "Sur devis", "Illimité", "Solution complète personnalisée"],
        ]

        for i, row_data in enumerate(pricing_data):
            for j, cell_data in enumerate(row_data):
                cell = table.cell(i + 1, j)
                cell.text = cell_data
                run = cell.paragraphs[0].runs[0]
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)

        doc.add_paragraph()  # Espace après le tableau

        self._add_word_subsection(doc, "Projections Financières")
        projections = [
            "• Objectif Année 1 : 500 clients, 150M FCFA ARR",
            "• Objectif Année 2 : 2,000 clients, 650M FCFA ARR",
            "• Objectif Année 3 : 5,000 clients, 1.8Md FCFA ARR",
            "• Rentabilité : Mois 18 avec 35% de marge nette",
        ]

        for projection in projections:
            self._add_word_highlight(doc, projection)

        # Section 7
        self._add_word_section_title(doc, "7. ÉQUIPE DE PROJET")

        # Profil du développeur principal
        self._add_word_subsection(
            doc, "Benny Joumessi Jason Isaac - Fondateur & Développeur Principal"
        )
        self._add_word_paragraph(
            doc,
            "Ingénieur en informatique spécialisé en intelligence artificielle, expert en développement "
            "d'applications cloud avec Azure AI Foundry. Expérience significative dans le développement "
            "de solutions SaaS pour PME. Employé chez GTA Company avec une expertise approfondie "
            "en architecture cloud-native et systèmes d'IA conversationnelle.",
        )

        self._add_word_subsection(doc, "Compétences Techniques")
        skills = [
            "• Développement Python avancé et architectures microservices",
            "• Expertise Azure AI Foundry, GPT-4o et services cloud Microsoft",
            "• Maîtrise des bases de données NoSQL (CosmosDB) et systèmes de cache",
            "• Développement d'interfaces utilisateur modernes avec Streamlit",
            "• Analytics et visualisation de données avec Plotly et Pandas",
        ]

        for skill in skills:
            self._add_word_highlight(doc, skill)

        self._add_word_subsection(doc, "Vision du Projet")
        self._add_word_paragraph(
            doc,
            "En tant que développeur unique de ce projet, l'objectif est de créer une solution "
            "robuste et scalable qui répond aux besoins spécifiques du marché camerounais. "
            "L'approche combine l'innovation technologique avec une compréhension profonde "
            "des défis locaux des PME en matière de facturation et de gestion financière.",
        )

    def _add_word_section_title(self, doc, text):
        """Ajoute un titre de section Word."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = None
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(12)

    def _add_word_subsection(self, doc, text):
        """Ajoute un sous-titre Word."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.font.bold = True
        p.paragraph_format.space_before = Pt(15)
        p.paragraph_format.space_after = Pt(8)

    def _add_word_paragraph(self, doc, text):
        """Ajoute un paragraphe normal Word."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _add_word_highlight(self, doc, text):
        """Ajoute un point important Word."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.bold = True
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.5)

    def _create_title_page(self):
        """Crée la page de titre PDF."""
        elements = []

        # Titre principal
        elements.append(Spacer(1, 1 * inch))
        elements.append(Paragraph("DOCUMENT DE PROJET", self.title_style))
        elements.append(Spacer(1, 0.5 * inch))

        # Nom du projet
        project_title = ParagraphStyle(
            "ProjectTitle",
            fontSize=18,
            fontName="Times-Bold",
            textColor=colors.HexColor("#048A81"),
            alignment=TA_CENTER,
            spaceAfter=20,
            leading=27,  # 1.5 * fontSize
        )
        elements.append(
            Paragraph("FacturIQ.ai<br/>La Facturation Intelligente", project_title)
        )
        elements.append(Spacer(1, 0.3 * inch))

        # Sous-titre
        subtitle_style = ParagraphStyle(
            "Subtitle",
            fontSize=14,
            fontName="Times-Roman",
            textColor=colors.HexColor("#666666"),
            alignment=TA_CENTER,
            leading=21,  # 1.5 * fontSize
        )
        elements.append(
            Paragraph(
                "Plateforme d'automatisation intelligente pour la gestion<br/>et génération de factures professionnelles",
                subtitle_style,
            )
        )
        elements.append(Spacer(1, 0.8 * inch))

        # Informations du document
        info_data = [
            ["Date:", datetime.now().strftime("%d/%m/%Y")],
            ["Version:", "2.0"],
            ["Statut:", "Production"],
            ["Technologie:", "Azure AI Foundry + GPT-4o"],
            ["Développeur:", "Benny Joumessi Jason Isaac"],
        ]

        info_table = Table(info_data, colWidths=[4 * cm, 6 * cm])
        info_table.setStyle(
            TableStyle(
                [
                    ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                    ("FONTNAME", (0, 0), (0, -1), "Times-Bold"),
                    ("FONTNAME", (1, 0), (1, -1), "Times-Roman"),
                    ("FONTSIZE", (0, 0), (-1, -1), 12),
                    ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#2E4057")),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )

        elements.append(info_table)
        elements.append(Spacer(1, 1 * inch))

        return elements

    def _create_table_of_contents(self):
        """Crée la table des matières PDF."""
        elements = []

        elements.append(Paragraph("TABLE DES MATIÈRES", self.section_style))
        elements.append(Spacer(1, 0.3 * inch))

        toc_data = [
            ["1. Présentation de la start-up et du projet", "3"],
            ["2. Problème à résoudre et opportunité de marché", "3"],
            ["3. Facteur innovant", "4"],
            ["4. Résumé du projet", "4"],
            ["5. Cible et bénéficiaires", "5"],
            ["6. Modèle économique", "5"],
            ["7. Équipe de projet", "6"],
        ]

        toc_table = Table(toc_data, colWidths=[14 * cm, 2 * cm])
        toc_table.setStyle(
            TableStyle(
                [
                    ("ALIGN", (0, 0), (0, -1), "LEFT"),
                    ("ALIGN", (1, 0), (1, -1), "RIGHT"),
                    ("FONTNAME", (0, 0), (-1, -1), "Times-Roman"),
                    ("FONTSIZE", (0, 0), (-1, -1), 12),
                    ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#333333")),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
                ]
            )
        )

        elements.append(toc_table)

        return elements

    def _create_main_content(self):
        """Crée le contenu principal du document PDF."""
        elements = []

        # Section 1: Présentation de la start-up et du projet
        elements.append(
            Paragraph("1. PRÉSENTATION DE LA START-UP ET DU PROJET", self.section_style)
        )

        elements.append(Paragraph("Notre Vision", self.subsection_style))
        elements.append(
            Paragraph(
                "FacturIQ.ai est une start-up technologique camerounaise spécialisée dans l'automatisation "
                "intelligente des processus de facturation pour les PME et entreprises de services. Notre mission "
                "est de révolutionner la gestion de facturation en combinant l'intelligence artificielle "
                "avancée avec les technologies cloud Microsoft Azure.",
                self.body_style,
            )
        )

        elements.append(Paragraph("Le Projet", self.subsection_style))
        elements.append(
            Paragraph(
                "FacturIQ.ai - La Facturation Intelligente est une plateforme complète qui automatise "
                "entièrement le cycle de vie des factures, de la génération à l'analyse des performances "
                "commerciales. Utilisant GPT-4o d'OpenAI via Azure AI Foundry, le système comprend le "
                "langage naturel et génère des factures professionnelles en quelques secondes.",
                self.body_style,
            )
        )

        # Section 2: Problème et opportunité
        elements.append(
            Paragraph(
                "2. PROBLÈME À RÉSOUDRE ET OPPORTUNITÉ DE MARCHÉ", self.section_style
            )
        )

        elements.append(Paragraph("Problématiques Identifiées", self.subsection_style))
        problems = [
            "• <b>Processus manuel chronophage</b> : Les PME passent en moyenne 5-8 heures par semaine sur la facturation",
            "• <b>Erreurs humaines coûteuses</b> : 23% des factures contiennent des erreurs, retardant les paiements",
            "• <b>Manque de visibilité</b> : Absence d'outils d'analyse pour optimiser la trésorerie",
            "• <b>Conformité réglementaire</b> : Difficultés à respecter les normes fiscales en évolution",
            "• <b>Processus de suivi défaillant</b> : Retards de paiement non anticipés",
        ]

        for problem in problems:
            elements.append(Paragraph(problem, self.highlight_style))

        elements.append(Paragraph("Opportunité de Marché", self.subsection_style))
        elements.append(
            Paragraph(
                "Le marché mondial des logiciels de facturation représente 18,7 milliards USD en 2024, "
                "avec une croissance annuelle de 15,4%. Au Cameroun et en Afrique centrale, les PME "
                "cherchent des solutions d'automatisation adaptées au contexte local. Notre approche "
                "IA-first nous positionne sur un segment à forte valeur ajoutée avec peu de concurrence directe.",
                self.body_style,
            )
        )

        # Section 3: Facteur innovant
        elements.append(Paragraph("3. FACTEUR INNOVANT", self.section_style))

        innovations = [
            '• <b>IA Conversationnelle</b> : Génération de factures par langage naturel ("Facture Acme Corp, 40h consulting, 75000 FCFA/h")',
            "• <b>Optimisation Avancée</b> : Système de cache intelligent offrant 23,464x d'amélioration des performances",
            "• <b>Analytics Prédictive</b> : Prévisions de trésorerie et scoring de risque client automatisés",
            "• <b>Architecture Résiliente</b> : Circuit breaker et rate limiting pour 99.9% de disponibilité",
            "• <b>Interface Unifiée</b> : Chat IA + dashboard analytics + génération rapide dans une seule plateforme",
        ]

        for innovation in innovations:
            elements.append(Paragraph(innovation, self.highlight_style))

        elements.append(
            Paragraph(
                "Notre différenciation clé réside dans l'intégration native de l'IA générative avec "
                "des capacités d'analyse business en temps réel, créant un assistant intelligent "
                "capable de comprendre le contexte métier camerounais et d'anticiper les besoins.",
                self.body_style,
            )
        )

        # Section 4: Résumé du projet
        elements.append(Paragraph("4. RÉSUMÉ DU PROJET", self.section_style))

        elements.append(Paragraph("Architecture Technique", self.subsection_style))
        elements.append(
            Paragraph(
                "• <b>Frontend</b> : Interface Streamlit responsive avec dashboard analytics interactif<br/>"
                "• <b>IA Core</b> : GPT-4o via Azure AI Foundry pour génération intelligente<br/>"
                "• <b>Backend</b> : Microservices Python avec pattern singleton et cache avancé<br/>"
                "• <b>Storage</b> : CosmosDB (données), Blob Storage (PDFs), AI Search (indexation)<br/>"
                "• <b>Analytics</b> : Plotly pour visualisations, algorithmes prédictifs personnalisés",
                self.highlight_style,
            )
        )

        elements.append(Paragraph("Fonctionnalités Clés", self.subsection_style))
        elements.append(
            Paragraph(
                "La plateforme offre une génération de factures en langage naturel, un système "
                "d'analytics business en temps réel, un cache intelligent réduisant les coûts API "
                "de 75%, et une architecture cloud-native garantissant scalabilité et résilience. "
                "L'interface unifiée combine chat IA, dashboard analytics et outils de gestion "
                "dans une expérience utilisateur moderne adaptée au marché camerounais.",
                self.body_style,
            )
        )

        # Section 5: Cible et bénéficiaires
        elements.append(Paragraph("5. CIBLE ET BÉNÉFICIAIRES", self.section_style))

        elements.append(Paragraph("Marché Primaire", self.subsection_style))
        targets = [
            "• <b>PME Services</b> (5-50 employés) : Consultants, agences, services professionnels au Cameroun",
            "• <b>Freelances Premium</b> : Consultants indépendants à forte facturation",
            "• <b>Start-ups B2B</b> : Entreprises SaaS en croissance cherchant l'automatisation",
            "• <b>Cabinets Spécialisés</b> : Comptables, avocats, architectes, ingénieurs",
        ]

        for target in targets:
            elements.append(Paragraph(target, self.highlight_style))

        elements.append(Paragraph("Bénéfices Quantifiés", self.subsection_style))
        elements.append(
            Paragraph(
                "• <b>Gain de temps</b> : 85% de réduction du temps de facturation<br/>"
                "• <b>Réduction erreurs</b> : 90% moins d'erreurs grâce à l'automatisation IA<br/>"
                "• <b>Amélioration trésorerie</b> : 25% d'accélération des encaissements<br/>"
                "• <b>Visibilité business</b> : Analytics prédictive pour optimisation continue",
                self.highlight_style,
            )
        )

        # Section 6: Modèle économique
        elements.append(Paragraph("6. MODÈLE ÉCONOMIQUE", self.section_style))

        elements.append(Paragraph("Stratégie de Monétisation", self.subsection_style))

        pricing_data = [
            ["Forfait", "Prix/mois", "Factures/mois", "Fonctionnalités"],
            ["Starter", "15,000 FCFA", "≤ 50", "Génération IA + Analytics de base"],
            ["Professional", "40,000 FCFA", "≤ 200", "+ Dashboard avancé + API"],
            [
                "Enterprise",
                "100,000 FCFA",
                "Illimité",
                "+ IA personnalisée + Support dédié",
            ],
            ["White Label", "Sur devis", "Illimité", "Solution complète personnalisée"],
        ]

        pricing_table = Table(pricing_data, colWidths=[3 * cm, 3 * cm, 3 * cm, 6 * cm])
        pricing_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#048A81")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
                    ("FONTNAME", (0, 1), (-1, -1), "Times-Roman"),
                    ("FONTSIZE", (0, 0), (-1, -1), 11),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                    ("BACKGROUND", (0, 1), (-1, -1), colors.HexColor("#F8F9FA")),
                    ("GRID", (0, 0), (-1, -1), 1, colors.HexColor("#E5E7EB")),
                ]
            )
        )

        elements.append(pricing_table)
        elements.append(Spacer(1, 0.2 * inch))

        elements.append(Paragraph("Projections Financières", self.subsection_style))
        elements.append(
            Paragraph(
                "• <b>Objectif Année 1</b> : 500 clients, 150M FCFA ARR<br/>"
                "• <b>Objectif Année 2</b> : 2,000 clients, 650M FCFA ARR<br/>"
                "• <b>Objectif Année 3</b> : 5,000 clients, 1.8Md FCFA ARR<br/>"
                "• <b>Rentabilité</b> : Mois 18 avec 35% de marge nette",
                self.highlight_style,
            )
        )

        # Section 7: Équipe de projet
        elements.append(Paragraph("7. ÉQUIPE DE PROJET", self.section_style))

        elements.append(
            Paragraph(
                "Benny Joumessi Jason Isaac - Fondateur & Développeur Principal",
                self.subsection_style,
            )
        )
        elements.append(
            Paragraph(
                "Ingénieur en informatique spécialisé en intelligence artificielle, expert en développement "
                "d'applications cloud avec Azure AI Foundry. Expérience significative dans le développement "
                "de solutions SaaS pour PME. Employé chez GTA Company avec une expertise approfondie "
                "en architecture cloud-native et systèmes d'IA conversationnelle.",
                self.body_style,
            )
        )

        elements.append(Paragraph("Compétences Techniques", self.subsection_style))
        skills = [
            "• <b>Développement Python avancé</b> et architectures microservices",
            "• <b>Expertise Azure AI Foundry</b>, GPT-4o et services cloud Microsoft",
            "• <b>Maîtrise des bases de données</b> NoSQL (CosmosDB) et systèmes de cache",
            "• <b>Développement d'interfaces utilisateur</b> modernes avec Streamlit",
            "• <b>Analytics et visualisation</b> de données avec Plotly et Pandas",
        ]

        for skill in skills:
            elements.append(Paragraph(skill, self.highlight_style))

        elements.append(Paragraph("Vision du Projet", self.subsection_style))
        elements.append(
            Paragraph(
                "En tant que développeur unique de ce projet, l'objectif est de créer une solution "
                "robuste et scalable qui répond aux besoins spécifiques du marché camerounais. "
                "L'approche combine l'innovation technologique avec une compréhension profonde "
                "des défis locaux des PME en matière de facturation et de gestion financière.",
                self.body_style,
            )
        )

        return elements


def main():
    """Fonction principale pour générer les documents."""
    print("🔄 Génération des documents de projet...")

    generator = ProjetDocumentGenerator()
    pdf_filename, word_filename = generator.generate_documents()

    # Vérification de la taille des fichiers
    pdf_size = os.path.getsize(pdf_filename) / (1024 * 1024)
    print(f"📁 Taille du fichier PDF: {pdf_size:.2f} MB")

    if pdf_size < 2.0:
        print("✅ Le fichier PDF respecte la limite de 2MB")
    else:
        print("⚠️ Attention: Le fichier PDF dépasse 2MB")

    if word_filename:
        word_size = os.path.getsize(word_filename) / (1024 * 1024)
        print(f"📁 Taille du fichier Word: {word_size:.2f} MB")

        if word_size < 2.0:
            print("✅ Le fichier Word respecte la limite de 2MB")
        else:
            print("⚠️ Attention: Le fichier Word dépasse 2MB")

    return pdf_filename, word_filename


if __name__ == "__main__":
    main()
